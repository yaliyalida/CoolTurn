# _*_ coding:utf-8 _*_
from docx import Document
import xlwt
import xlrd
from xlutils.copy import copy
import sys
import os


#用户上传的excel表头文件

startExcel = r'D:\Python\Cool_Turn\学生信息表表头.xlsx'
#最后生成的excel，这个库只能保存xls格式的文件
endExcel = r'D:\Python\Cool_Turn\学生信息表.xls'
#模板表
templete = r'D:\Python\Cool_Turn\模板表.docx'
#word所在的文件夹
wordDir = r'D:\Python\Cool_Turn\数据'
#模板表中每个字段对应的位置，键是字段，值是所在的位置
dict1 = {}

#判断是否是英文
def isEnglish(checkStr):
    for ch in checkStr.encode('utf-8').decode('utf-8'):
        if u'\u4e00' <= ch <= u'\u9fff':
            return False
    return True


#读取模板表
def readTemplate():

    document = Document(templete.encode('utf-8').decode('utf-8'))
    tempTable = document.tables
    table = tempTable[0]

    rowList = table.rows
    columnList = table.columns
    rowLength = len(rowList)
    columnLength = len(columnList)

    for rowIndex in range(rowLength):
        for columnIndex in range(columnLength):
            cell = table.cell(rowIndex,columnIndex)
            if isEnglish(cell.text):
                dict1.setdefault(cell.text,[rowIndex,columnIndex])

#读入的表
re = xlrd.open_workbook(startExcel.encode('utf-8').decode("utf-8"))
#通过复制读入的表来生成写入的表
we = copy(re)

#写第一页的sheet
def writeFirstSheet1(table, row):

    sheet = we.get_sheet(0)
    #将字段对应的值填到sheet1dict中
    sheet1dict = {}
    for key in dict1:
        tempList = dict1[key]
        for index in range(0,1):
            x = tempList[index]
            y = tempList[index+1]
            sheet1dict.setdefault(key,table.cell(x,y).text)


    #读取第一个sheet
    tempSheet = re.sheet_by_index(0)
    #读取第一个sheet中的第二行
    list1 = tempSheet.row_values(1)
    for excelIndex in range(len(list1)):
        for key in sheet1dict:
            if list1[excelIndex] == key:
                #将sheet1dict中的内容写入excel的sheet中
                sheet.write(row, excelIndex, sheet1dict[key])

#将word中数据写入excel
def writeExcel(wordName, row):
    document = Document(wordName)
    tempTable = document.tables
    table = tempTable[0]

    #一个excel一般有好几个sheet（即页数），所以单独写一个函数
    writeFirstSheet1(table, row)

    we.save(endExcel.encode('utf-8').decode("utf-8"))

if __name__ == "__main__":

    readTemplate()
    docFiles = os.listdir(wordDir.encode('utf-8').decode("utf-8"))
    # 开始数据的行数
    row = 1
    for doc in docFiles:
        #输出文件名
        print(doc.encode('utf-8').decode("utf-8"))
        try:
            row += 1
            writeExcel(wordDir + '\\' + doc.encode('utf-8').decode("utf-8"), row)
        except Exception as e:
            print(e)
import docx
import re
import pandas as pd
import os


def get_filelist(dir_path):
    filelist = []
    for file in os.scandir(dir_path):
        filelist.append(file.path)
    return filelist

def read_tplt(tplt_word):
    dic_fill = {}  # 记录信息位置与要填写位置的映射字典
    pattern = '{\d+}'  # 设定模式
    document = docx.Document(tplt_word)
    tbobj_list = document.tables  # 返回列表类型
    tbobj = tbobj_list[0]  # 文件中第一个表格对象
    row_num = len(tbobj.rows)
    col_num = len(tbobj.columns)
    for row_index in range(row_num):
        for col_index in range(col_num):
            cell = tbobj.cell(row_index,col_index)  # 单元格
            search_obj = re.search(pattern,cell.text)
            if search_obj:  # 查找不到则为False
                dic_fill.setdefault(search_obj.group()[1:-1],(row_index,col_index))  # 键为去除{}提取其中数字,值为cell位置
    # dic_fill = dict(zip(dic_fill.values(),dic_fill.keys()))  # zip打包为元组,这里使字典键值互换
    return dic_fill

def read_format_excel(format_excel):
    df = pd.read_excel(format_excel)
    return df

def add_data(filelist,dic_fill,df):
    dic_length = len(dic_fill)
    for file in filelist:
        ls_data = ['NaN']*dic_length  # 初始化行数据
        document = docx.Document(file)
        tbobj_list = document.tables  # 返回列表类型
        tbobj = tbobj_list[0]  # 文件中第一个表格对象
        for key,value in dic_fill.items():
            row,col = value[0],value[1]
            ls_data[int(key)-1] = tbobj.cell(row,col).text
        df.loc[len(df)] = ls_data
    return df

def write_excel(df,save_path='result.xlsx'):
    try:
        df.to_excel(save_path,index=False)
        print("当前路径是{}".format(os.getcwd()))
        print("{} 存储成功".format(save_path))
    except Exception as err:
        print(err)
        print("存储失败")

def main():
    print("请输入word模板文件路径:")
    tplt_word = input().replace('"','')  # word模板
    print("请输入excel模板文件路径:")
    format_excel = input().replace('"','')  # excel模板
    print("请输入要进行汇总的word文件夹路径:")
    dir_path = input().replace('"','')  # 数据文件夹
    # 模板表中每个字段对应的位置，键是字段，值是所在的位置
    try:
        dic_fill = read_tplt(tplt_word)
        df = read_format_excel(format_excel)
        filelist = get_filelist(dir_path)
        df = add_data(filelist, dic_fill, df)
        write_excel(df)
    except IndexError:
        print("请检查word模板")
    except Exception as err:
        print(err)
        print("请检查输入的文件以及文件夹路径")


if __name__ == "__main__":
    tplt_word = 'template_word.docx'  # word模板
    format_excel = 'template_excel.xlsx'  # excel模板
    dir_path = 'example'  # 数据文件夹
    # 模板表中每个字段对应的位置，键是字段，值是所在的位置
    try:
        dic_fill = read_tplt(tplt_word)
        df = read_format_excel(format_excel)
        filelist = get_filelist(dir_path)
        df = add_data(filelist,dic_fill,df)
        write_excel(df)
    except IndexError:
        print("请检查word模板")
    except Exception as err:
        print(err)
        print("请检查设置的文件以及文件夹路径")
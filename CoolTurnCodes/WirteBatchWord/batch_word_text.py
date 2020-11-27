import docx
import os
from docx.oxml.ns import qn


def read_text(formatpath):
    '''
    获取模板文件字符串,段落之间用换行符链接
    :param formatpath: 模板文件路径
    :return: 模板文件字符内容
    '''
    content = ''
    try:
        fileobj = docx.Document(formatpath)  # 模板文件对象
        for para in fileobj.paragraphs:
            content += para.text+'\n'
    except Exception as err:
        print(err)
    return content

def batch_write(filenum,fieldnum):
    '''
    根据用户指定的数量批量填写属性值
    :param filenum: 用户指定的批量填写的文件数目
    :param fieldnum: 每一份word中需要填写的字段数目
    :return: 返回类型是列表的列表,存储的元素是以同一字段每一个文件内容为元素的列表
    '''
    fields = []
    optional = '是否YyNn'
    sure = '是Yy'
    for i in range(fieldnum):  # 逐一遍历字段
        files = []
        for j in range(filenum):  # 逐一遍历文件
            enter = input("请输入第{}个字段,第{}个文件的内容:\n"\
                               .format(i+1,j+1))
            if j==0:  # 第一次输入
                while True:
                    try:
                        choice = input("需要设置该字段为公共字段吗?(是/否)(y/n):\n")
                        if choice not in optional:  # 如果输入不在可选字符范围
                            raise ValueError("需输入'是'、'否'、'y'、'n'中的一个字符")
                        break
                    except Exception as err:
                        print("输入不符合要求:{}\n请重新输入".format(repr(err)))
                if choice in sure:  # 如果设置为公共字段
                    print("已将该公共字段值批量填写到每一份word文件中")
                    files.append(enter)
                    break

            if enter == '':  # 后续还可通过不输入内容来设置公共字段
                while True:
                    try:
                        choice = input("检测到此次输入为空,是否需要重新输入？(是/否)(y/n):\n")
                        if choice not in optional:  # 如果输入不在可选字符范围
                            raise ValueError("需输入'是'、'否'、'y'、'n'中的一个字符")
                        break
                    except Exception as err:
                        print("输入不符合要求:{}\n请重新输入".format(repr(err)))
                if choice in sure:  # 如果需要重新输入
                    enter = input("请重新输入第{}个字段,第{}个文件的内容:\n" \
                                  .format(i + 1, j + 1))

            files.append(enter)
        fields.append(files)
        prompt = "第{}字段的内容填写完成".format(i+1)
        print("{:-^30}".format(prompt))

    return fields

def write_words(template,specific_chrt,fields,filenum,name_rules=0,save_dir='WordsResult'):
    '''
    将以嵌套列表形式存储的文件内容写入文件,先确定好要写入的文件内容,尽量减少文件IO次数
    :param template: 模板字符串
    :param specific_chrt: 采用指定字符标记填写位置
    :param fields: 列表的列表,存储的元素是以同一字段每一个文件内容为元素的列表
    :return:
    '''
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)
    for i in range(filenum):  # 遍历每一个文件
        content = template  # 初始化要写入每一个文件的内容
        document = docx.Document()  # 创建一个空白文档对象
        # 设置文档字体
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        for j in range(len(fields)):  # 遍历每一个字段
            # 每次将一个指定字符specific_chrt替换为该文件对应的字段内容
            if len(fields[j]) < filenum:  # 如果该字段填写数目小于文件数目则判定为公共字段
                # 公共字段的值以第一个文件的值为准
                content = content.replace(specific_chrt,fields[j][0],1)
            else:  # 否则是非公共字段
                content = content.replace(specific_chrt,fields[j][i],1)
        # 添加段落,文本包含换行符
        document.add_paragraph(content)
        # 保存.docx文档
        if name_rules == 0:  # 如果命名规律采用默认设置
            filename = str(i+1) + '.docx'
        else:  # 用户采用自定义命名规律
            if len(fields[(name_rules)-1]) < filenum:
                print("不支持采用公共字段值命名文件,转而采用默认设置")
                filename = str(i+1) + '.docx'
            else:
                filename = fields[(name_rules)-1][i] + '.docx'
        save_path = save_dir + '\\' + filename
        try:  # 涉及文件IO,进行异常处理
            document.save(save_path)
            # 输出文件存储路径的提示信息
            current_path = os.getcwd()  # 获得当前路径
            print("当前路径是{}".format(current_path))
            print("{} 存储成功".format(save_path))
        except Exception as err:
            print(err)
            print("文件存储失败")

def enter_filenum():
    while True:  # 直到用户输入正确跳出
        try:
            filenum = int(input("请输入想批量写入的文件数量:\n"))
            if filenum <= 0:
                raise ValueError("输入有误,需为正整数")
            break
        except Exception as err:
            print("输入不符合要求:{}\n请重新输入".format(repr(err)))
    return filenum

def enter_name_rules(fields_num):
    while True:
        try:
            name_rules = int(input\
                ("文件命名采用默认命名/第几个字段命名(0/整数{}~{}):\n".format(1,fields_num)))
            if name_rules < 0:
                raise ValueError("输入有误,不能是负数")
            break
        except Exception as err:
            print("输入不符合要求:{}\n请重新输入".format(repr(err)))
    return name_rules

def main(formatpath):
    template = read_text(formatpath)  # 模板文件内容
    print("模板内容为:\n{}".format(template))
    specific_chrt = "{}"  # 采用指定字符{}标记填写位置
    fieldnum = template.count(specific_chrt)  # 需要填写的字段的数量
    print("每份word中需要填写的字段数量:{}".format(fieldnum))
    filenum = enter_filenum()  # 由用户指定需要批量填写的文件数量
    fields = batch_write(filenum, fieldnum)  # 根据文件数量和字段数量批量填写内容
    name_rules = enter_name_rules(len(fields))  # 由用户指定命名规则
    write_words(template, specific_chrt, fields, filenum, name_rules)  # 将批量填写的内容逐一写入文件

# 字符串方法速度上比正则表达式方法更快 content
if __name__ == "__main__":
    formatpath = 'template_text.docx'  # 模板文件路径
    main(formatpath)




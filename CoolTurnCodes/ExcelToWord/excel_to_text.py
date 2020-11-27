from openpyxl import load_workbook
import docx
from docx.oxml.ns import qn
import re
import os


def get_tplt(formatpath):
    '''
    获取模板文件字符串,段落之间用换行符链接
    :param formatpath: 模板文件路径
    :return: 模板文件字符内容
    '''
    content = ''
    try:
        fileobj = docx.Document(formatpath)  # 模板文件对象
        for para in fileobj.paragraphs:
            content += para.text + '\n'
    except Exception as err:
        print(err)
    return content

# 设置文档字体
def set_font(document):
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

def get_ws(file_path):
    # 读取excel xlsx文件
    wb = load_workbook(file_path)  # 打开现有工作表
    ws = wb.active  # 默认对第一张工作表进行操作
    return ws

def get_title(ws):
    title = []
    for col_index in range(ws.max_column):
        title.append(ws.cell(row=1, column=col_index+1).value)
    return title

def print_title(title):
    print("表头字段如下:")
    for t in title:
        print(t,end=" ")
    print()

def enter_choice():
    optional = '是否YyNn'
    sure = '是Yy'
    while True:
        try:
            choice = input("是否采用与字段内容无关的数值递增的文件命名方式?(是/否)(y/n):\n")
            if choice not in optional:  # 如果输入不在可选字符范围
                raise ValueError("需输入'是'、'否'、'y'、'n'中的一个字符")
            break
        except Exception as err:
            print("输入不符合要求:{}\n请重新输入".format(repr(err)))
    if choice in sure:
        return True
    else:
        return False

def enter_name_rules(title):
    while True:
        try:
            print("请输入命名字段")
            name_title = input()
            if name_title not in title:
                raise ValueError("请原样输入表头中的一个字段")
            name_rules = title.index(name_title)
            break
        except Exception as err:
            print(err)
    return name_rules

# 转换成word文本
def excel_to_text(ws, template, name_rules, default_value, save_dir='ExcelToWordResult'):

    if not os.path.exists(save_dir):
        os.makedirs(save_dir)
    # 获取excel行数
    row_num = ws.max_row
    # 匹配模板字符
    tplt_list = re.findall(r'{\d+}', template)  # 以匹配到的元素列表返回

    # 跳过表头

    # 写入word文件
    for row_index in range(2,row_num+1):
        # 初始化写入内容
        content = template
        for tplt in tplt_list:
            fill_index = int(tplt[1:-1])  # 找到要填入的内容对应excel表格的单元格,提取{}中的数值
            try:
                fill_value = str(ws.cell(row=row_index, column=fill_index).value)  # 找到要替换的内容
                fill_value = '?' if fill_value=='None' else fill_value  # 如果下标越界,则该位置内容为None
                content = content.replace(tplt, fill_value)  # 替换
            except Exception as err:
                print(err)
                print("Please check")

        # 创建word文档
        document = docx.Document()
        # 设置文档字体
        set_font(document)
        # 向文档中写入内容
        document.add_paragraph(content)

        # 文件名
        if name_rules == default_value:  # 如果采用默认命名(数字递增)
            filename = str(row_index - 1) + '.docx'
        # 文件名
        else:
            filename = str(ws.cell(row=row_index, column=name_rules + 1).value) + '.docx'

        # 保存文件
        save_path = save_dir + '\\' + filename
        try:  # 涉及文件IO,进行异常处理
            document.save(save_path)
            # 输出文件存储路径的提示信息
            current_path = os.getcwd()  # 获得当前路径
            print("当前路径是{}".format(current_path))
            print("{} 存储成功".format(save_path))
        except Exception as err:
            print(err)
            print("文件存储失败")

def main(file_path,formatpath):
    ws = get_ws(file_path)  # 获取工作簿对象
    template = get_tplt(formatpath)  # 模板字符串
    title = get_title(ws)  # 获取其表头字段
    print_title(title)
    choice = enter_choice()  # 由用户指定是否采用数值递增命名
    default_value = -1  # 命名方式缺省值
    if choice:
        name_rules = default_value
    else:
        name_rules = enter_name_rules(title)
    excel_to_text(ws, template, name_rules, default_value)  # 以文本形式写入批量word

if __name__ == "__main__":
    # 用户指定
    file_path = '学生信息表.xlsx'  # 要处理的excel文件路径
    formatpath = 'template_text.docx'
    main(file_path, formatpath)

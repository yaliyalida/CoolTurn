import docx
from docx.oxml.ns import qn

def read_table(filename):
    doc = docx.Document(filename)
    tables = doc.tables
    print("There are {} tables".format(len(tables)))
    paras = doc.paragraphs
    for para in paras:
        print(para.text)
    print("There are {} paragraphs".format(len(paras)))
    ls_tables = []
    for table in tables:
        ls_table = []
        for row in table.rows:
            ls_row = []
            for cell in row.cells:
                ls_row.append(cell.text)
            ls_table.append(ls_row)
        ls_tables.append(ls_table)
    return ls_tables

def write_table(ls_tables):
    savename = 'test.docx'
    doc = docx.Document()
    set_font(doc)
    for table in ls_tables:
        rownums = len(table)
        colnums = len(table[0])
        tbobj = doc.add_table(rownums,colnums,style="Table Grid")
        doc.add_paragraph()
        for i in range(rownums):
            for j in range(colnums):
                tbobj.cell(i,j).text = table[i][j]
    doc.save(savename)
    print("finish")

# 设置文档字体
def set_font(doc):
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 计数填写数量
def count_fill(ls_template,specific_chrt):
    count = 0
    for table in ls_template:
        for row in table:
            count += row.count(specific_chrt)
    return count

if __name__ == "__main__":
    formatname = 'template_table.docx'
    ls_template = read_table(filename)
    print("模板内容为:\n{}".format(ls_template))
    specific_chrt = "{}"  # 采用指定字符{}标记填写位置
    count = count_fill(ls_template,specific_chrt)
    print("每份word中需要填写的字段数量:{}".format(fieldnum))
    filenum = 0  # 初始化用户想要批量写入的文件数量

    while True:  # 直到用户输入正确跳出
        try:
            filenum = int(input("请输入想批量写入的文件数量:\n"))
            if filenum <= 0:
                raise ValueError("输入有误,需为正整数")
            break
        except Exception as err:
            print("输入不符合要求:{}\n请重新输入".format(repr(err)))

    fields = batch_write(filenum, fieldnum)









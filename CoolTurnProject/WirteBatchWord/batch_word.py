import docx
from docx.oxml.ns import qn
import os


def read_text(formatpath):
    '''
    获取模板文件字符串,段落之间用换行符链接
    :param formatpath: 模板文件路径
    :return: 模板文件字符内容
    '''
    content = ''
    try:
        fileobj = docx.Document(formatpath)  # 模板文件对象
        for para in fileobj.paragraphs:
            content += para.text+'\n'
    except Exception as err:
        print(err)
    return content

def read_table(filename):
    doc = docx.Document(filename)
    tables = doc.tables
    print("There are {} tables".format(len(tables)))
    ls_tables = []
    for table in tables:
        ls_table = []
        for row in table.rows:
            ls_row = []
            for cell in row.cells:
                ls_row.append(cell.text)
            ls_table.append(ls_row)
        ls_tables.append(ls_table)
    return ls_tables

# 设置文档字体
def set_font(document):
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 计数填写数量
def count_fill(ls_template,specific_chrt):
    '''
    计算模板表格中需要填写的字段数量
    :param ls_template: 储存模板信息的多维列表
    :param specific_chrt: 用来标记填充的特殊字符
    :return: 模板表格中需要填写的字段数量
    '''
    count = 0
    for table in ls_template:
        for row in table:
            count += row.count(specific_chrt)
    return count

def batch_write(filenum,fieldnum):
    '''
    根据用户指定的数量批量填写属性值
    :param filenum: 用户指定的批量填写的文件数目
    :param fieldnum: 每一份word中需要填写的字段数目
    :return: 返回类型是列表的列表,存储的元素是以同一字段每一个文件内容为元素的列表
    '''
    fields = []
    optional = '是否YyNn'
    sure = '是Yy'
    for i in range(fieldnum):  # 逐一遍历字段
        files = []
        for j in range(filenum):  # 逐一遍历文件
            enter = input("请输入第{}个字段,第{}个文件的内容:\n"\
                               .format(i+1,j+1))
            if j==0:  # 第一次输入
                while True:
                    try:
                        choice = input("需要设置该字段为公共字段吗?(是/否)(y/n):\n")
                        if choice not in optional:  # 如果输入不在可选字符范围
                            raise ValueError("需输入'是'、'否'、'y'、'n'中的一个字符")
                        break
                    except Exception as err:
                        print("输入不符合要求:{}\n请重新输入".format(repr(err)))
                if choice in sure:  # 如果设置为公共字段
                    print("已将该公共字段值批量填写到每一份word文件中")
                    files.append(enter)
                    break

            if enter == '':
                while True:
                    try:
                        choice = input("检测到此次输入为空,是否需要重新输入？(是/否)(y/n):\n")
                        if choice not in optional:  # 如果输入不在可选字符范围
                            raise ValueError("需输入'是'、'否'、'y'、'n'中的一个字符")
                        break
                    except Exception as err:
                        print("输入不符合要求:{}\n请重新输入".format(repr(err)))
                if choice in sure:  # 如果需要重新输入
                    enter = input("请重新输入第{}个字段,第{}个文件的内容:\n" \
                                  .format(i + 1, j + 1))

            files.append(enter)
        fields.append(files)
        prompt = "第{}字段的内容填写完成".format(i+1)
        print("{:-^30}".format(prompt))
    return fields

def write_text_words(template,specific_chrt,fields,filenum,name_rules=0,save_dir='WordsResult'):
    '''
    将以嵌套列表形式存储的文件内容写入文件,先确定好要写入的文件内容,尽量减少文件IO次数
    :param template: 模板字符串
    :param specific_chrt: 采用指定字符标记填写位置
    :param fields: 列表的列表,存储的元素是以同一字段每一个文件内容为元素的列表
    :return:
    '''
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)
    for i in range(filenum):  # 遍历每一个文件
        content = template  # 初始化要写入每一个文件的内容
        document = docx.Document()  # 创建一个空白文档对象
        # 设置文档字体
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        for j in range(len(fields)):  # 遍历每一个字段
            # 每次将一个指定字符specific_chrt替换为该文件对应的字段内容
            if len(fields[j]) < filenum:  # 如果该字段填写数目小于文件数目则判定为公共字段
                # 公共字段的值以第一个文件的值为准
                content = content.replace(specific_chrt,fields[j][0],1)
            else:  # 否则是非公共字段
                content = content.replace(specific_chrt,fields[j][i],1)
        # 添加段落,文本包含换行符
        document.add_paragraph(content)
        # 保存.docx文档
        if name_rules == 0:  # 如果命名规律采用默认设置
            filename = str(i+1) + '.docx'
        else:  # 用户采用自定义命名规律
            if len(fields[(name_rules)-1]) < filenum:
                print("不支持采用公共字段值命名文件,转而采用默认设置")
                filename = str(i+1) + '.docx'
            else:
                filename = fields[(name_rules)-1][i] + '.docx'
        save_path = save_dir + '\\' + filename
        try:  # 涉及文件IO,进行异常处理
            document.save(save_path)
            # 输出文件存储路径的提示信息
            current_path = os.getcwd()  # 获得当前路径
            print("当前路径是{}".format(current_path))
            print("{} 存储成功".format(save_path))
        except Exception as err:
            print(err)
            print("文件存储失败")

def write_table_words(ls_template,specific_chrt,fields,filenum,name_rules=0,save_dir='WordsResult'):
    '''

    :param ls_template: 存储模板信息的嵌套列表
    :param specific_chrt: 标记填入位置的特殊字符
    :param fields: 存储填入内容的嵌套列表
    :param filenum: 指定的批量写入数量
    :param name_rules: 命名规则,默认或者指定字段
    :return:
    '''
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)
    for file_index in range(filenum):  # 遍历每一个文件,减少IO
        field_index = 0
        document = docx.Document()  # 创建一个空白文档对象
        # 设置文档字体
        set_font(document)
        for table in ls_template:
            rownums = len(table)
            colnums = len(table[0])
            tbobj = document.add_table(rows=rownums, cols=colnums, style="Table Grid")
            for i in range(rownums):
                for j in range(colnums):
                    if table[i][j] == specific_chrt:  # 如果该位置是需要填充位置
                        if len(fields[field_index]) < filenum:  # 判定是公共字段
                            tbobj.cell(i, j).text = fields[field_index][0]
                        else:  # 否则是非公共字段
                            tbobj.cell(i, j).text = fields[field_index][file_index]
                        field_index += 1
                    else:
                        tbobj.cell(i,j).text = table[i][j]
            document.add_paragraph()  # 表之间以段落分隔

        # 保存.docx文档
        if name_rules == 0:  # 如果命名规律采用默认设置
            filename = str(file_index+1) + '.docx'
        else:  # 用户采用自定义命名规律
            if len(fields[(name_rules)-1]) < filenum:
                print("不支持采用公共字段值命名文件,转而采用默认设置")
                filename = str(file_index+1) + '.docx'
            else:
                filename = fields[(name_rules)-1][file_index] + '.docx'
        save_path = save_dir + '\\' + filename
        try:  # 涉及文件IO,进行异常处理
            document.save(save_path)
            # 输出文件存储路径的提示信息
            current_path = os.getcwd()  # 获得当前路径
            print("当前路径是{}".format(current_path))
            print("{} 存储成功".format(save_path))
        except Exception as err:
            print(err)
            print("文件存储失败")

def enter_filenum():
    while True:  # 直到用户输入正确跳出
        try:
            filenum = int(input("请输入想批量写入的文件数量:\n"))
            if filenum <= 0:
                raise ValueError("输入有误,需为正整数")
            break
        except Exception as err:
            print("输入不符合要求:{}\n请重新输入".format(repr(err)))
    return filenum

def enter_name_rules(fields_num):
    while True:
        try:
            name_rules = int(input\
                ("文件命名采用默认命名/第几个字段命名(0/整数{}~{}):\n".format(1,fields_num)))
            if name_rules < 0:
                raise ValueError("输入有误,不能是负数")
            break
        except Exception as err:
            print("输入不符合要求:{}\n请重新输入".format(repr(err)))
    return name_rules

def table_word(formatname):
    ls_template = read_table(formatname)  # 模板文件内容,以列表形式存储
    print("模板内容为:\n{}".format(ls_template))
    specific_chrt = "{}"  # 采用指定字符{}标记填写位置
    fieldnum = count_fill(ls_template, specific_chrt)  # 计算模板中需要填写的字段数量
    print("每份word中需要填写的字段数量:{}".format(fieldnum))
    filenum = enter_filenum()  # 由用户指定批量文件数量
    fields = batch_write(filenum, fieldnum)  # 根据文件数量和字段数量批量填写内容
    name_rules = enter_name_rules(len(fields))  # 由用户指定命名规则为默认或者某个字段值
    write_table_words(ls_template, specific_chrt, fields, filenum, name_rules)  # 将批量填写的内容逐一写入文件

def text_word(formatpath):
    template = read_text(formatpath)  # 模板文件内容
    print("模板内容为:\n{}".format(template))
    specific_chrt = "{}"  # 采用指定字符{}标记填写位置
    fieldnum = template.count(specific_chrt)  # 需要填写的字段的数量
    print("每份word中需要填写的字段数量:{}".format(fieldnum))
    filenum = enter_filenum()  # 由用户指定需要批量填写的文件数量
    fields = batch_write(filenum, fieldnum)  # 根据文件数量和字段数量批量填写内容
    name_rules = enter_name_rules(len(fields))  # 由用户指定命名规则
    write_text_words(template, specific_chrt, fields, filenum, name_rules)  # 将批量填写的内容逐一写入文件

def enter_choice():
    optional = ('1', '2')
    while True:
        try:
            print("功能选择:\n1、批量写入文本形式\n2、批量写入表格形式\n请选择:(1/2):")
            choice = input()
            if choice not in optional:
                raise ValueError('请输入"1"或"2"')
            break
        except Exception as err:
            print(err)
    return choice

def main():
    choice = enter_choice()
    if choice == '1':  # 批量写入文本形式
        while True:
            print("请输入模板文件绝对路径:")
            formatpath = input().replace('"','')
            try:
                text_word(formatpath)
                break
            except Exception as err:
                print(err)
                print("请检查文件路径是否正确")
    elif choice == '2':  # 批量写入表格形式
        while True:
            print("请输入模板文件绝对路径:")
            formatpath = input().replace('"','')
            try:
                table_word(formatpath)
                break
            except Exception as err:
                print(err)
                print("请检查文件路径是否正确")


if __name__ == "__main__":
    main()





